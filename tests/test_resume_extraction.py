"""Resume text extraction — Phase 8 (docs/product-plan.md). Extraction
only; candidate_analysis.py's own tests already cover what happens to
the extracted text once it reaches the model."""

import io

import pytest
from docx import Document
from pypdf import PdfWriter

from gtm_sourcing_agent import resume_extraction


def test_extract_text_from_txt():
    text = resume_extraction.extract_text("resume.txt", b"Jane Doe, Enterprise AE")
    assert text == "Jane Doe, Enterprise AE"


def test_extract_text_from_docx():
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Enterprise Account Executive")
    buffer = io.BytesIO()
    doc.save(buffer)

    text = resume_extraction.extract_text("resume.docx", buffer.getvalue())

    assert "Jane Doe" in text
    assert "Enterprise Account Executive" in text


def test_extract_text_from_blank_pdf_returns_empty_string():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)

    text = resume_extraction.extract_text("resume.pdf", buffer.getvalue())

    assert text == ""


def test_extract_text_rejects_unsupported_extension():
    with pytest.raises(ValueError, match="unsupported file type"):
        resume_extraction.extract_text("resume.rtf", b"whatever")
